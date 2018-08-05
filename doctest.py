def main():
    import docx
    doc1 = docx.Document("patent patent Alice.docx")
    doc2 = docx.Document()
    fulltext = []
    for para in doc1.paragraphs:
        fulltext.append(para.text)
    print(fulltext)
    for para in fulltext:
        doc2.add_paragraph(para)
        doc2.save("doc2.docx")
    print("MOM KNOWS EVERYTHING NOW")

main()
